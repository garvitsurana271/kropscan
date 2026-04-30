"""Generate pitch docx matching the Google Doc format exactly"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(1)
    s.bottom_margin = Inches(1)
    s.left_margin = Inches(1.2)
    s.right_margin = Inches(1.2)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

GREEN = RGBColor(56, 118, 29)  # Google Docs dark green
MAGENTA = RGBColor(194, 24, 91)
DARK = RGBColor(0, 0, 0)
GRAY = RGBColor(100, 100, 100)


def add_divider():
    p = doc.add_paragraph()
    # Add a horizontal line using bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)


def add_section_header(text, color):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.color.rgb = color
    r.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)


def add_spoken(text, bold_words=None):
    """Add spoken text with optional bold keywords"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(6)

    if not bold_words:
        r = p.add_run(text)
        r.font.size = Pt(12)
        r.font.color.rgb = DARK
        return

    # Split text around bold words and rebuild with formatting
    remaining = text
    for bw in bold_words:
        if bw in remaining:
            before, after = remaining.split(bw, 1)
            if before:
                r = p.add_run(before)
                r.font.size = Pt(12)
                r.font.color.rgb = DARK
            r = p.add_run(bw)
            r.font.size = Pt(12)
            r.font.color.rgb = DARK
            r.bold = True
            remaining = after

    if remaining:
        r = p.add_run(remaining)
        r.font.size = Pt(12)
        r.font.color.rgb = DARK


def add_direction(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = GREEN
    r.italic = True
    p.paragraph_format.space_after = Pt(6)


# ═══════════════════════════════════════════
# THE PITCH
# ═══════════════════════════════════════════

add_section_header('DEMO [0:00\u20140:15] (15 secs) \u2013 Garvit', GREEN)

add_spoken('Can I diagnose a crop disease with no internet? Let\u2019s find out.')

add_direction('Upload wheat_brown_rust.jpg. Wait for animation + chime. ~3 seconds.')

add_spoken(
    'Wheat Brown Rust. Propiconazole. Eight hundred rupees per acre. And it only took three seconds. That too without any internet. Everything you just saw ran on this laptop \u2014 no server, no cloud, no signal.',
    ['three seconds', 'any internet']
)

add_divider()

add_section_header('THE PROBLEM [0:15\u20140:35] (20 secs) \u2013 Garvit', GREEN)

add_spoken(
    'BUT a farmer in India waits five days for that same answer, and that too, if they can find an expert at all. One agro officer for every eleven hundred farmers. Ninety thousand crore lost to disease every single year. Not because it can\u2019t be treated, but because farmers can\u2019t get a diagnosis in time.',
    ['BUT', 'and that too', 'but']
)

add_divider()

add_section_header('WHY US [0:35\u20140:50] (15 secs) \u2013 Zaviaa', MAGENTA)

add_spoken(
    'I\u2019m Zaviaa. This is Garvit. We\u2019re from Assam, where farming is all around us. And we\u2019ve watched this happen, every monsoon, every season. A farmer spots a problem, goes to a dealer, gets the wrong advice, and the crop dies. We were done watching.'
)

add_divider()

add_section_header('THE TECH [0:50\u20141:20] (30 secs) \u2013 Garvit', GREEN)

add_spoken(
    'So we built KropScan. Every app out there \u2014 Plantix, CropIn \u2014 needs the internet. Forty percent of rural India doesn\u2019t have that. So we compiled Google\u2019s EfficientNetV2 \u2014 a full neural network \u2014 to run inside the browser using WebAssembly. No server. No GPU. Just the phone. That\u2019s what you just saw. And when there IS internet? The AI goes one step further \u2014 if it\u2019s not fully confident, it quietly verifies with Gemma 3 cloud. The farmer just gets the right answer \u2014 and the badge on top tells them exactly how.',
    ['needs the internet', 'EfficientNetV2', 'WebAssembly', 'No server. No GPU. Just the phone.', 'IS internet', 'Gemma 3']
)

add_direction('Points at the result on the projected screen.')

add_spoken(
    'Trained on almost 2 lakh real field images \u2014 20 crops, 156 diseases.',
    ['2 lakh', '20 crops', '156 diseases']
)

add_divider()

add_section_header('LANGUAGES + FEATURES [1:20\u20141:32] (12 secs) \u2013 Garvit', GREEN)

add_direction('Turns off airplane mode. Clicks Hindi in the language switcher. The whole UI changes on screen.')

add_spoken(
    'Along with this \u2014 eleven languages, one-tap treatment purchase, plus live mandi prices, weather, crop planner, and an AI chatbot. Everything a farmer needs, in one place.',
    ['eleven languages']
)

add_divider()

add_section_header('BUSINESS & SCALE [1:32\u20141:55] (23 secs) \u2013 Zaviaa', MAGENTA)

add_spoken('Now this is live and deployed \u2014 you can scan a leaf on your phone after we sit down.')

add_direction('Beat.')

add_spoken(
    'Free for farmers \u2014 full diagnosis, treatment, 11 languages.',
    ['Free for farmers']
)

add_direction('Beat.')

add_spoken(
    'A farmer should never see a lock icon over their medicine.',
    ['never see a lock icon']
)

add_spoken(
    'There\u2019s a Pro plan at 49 rupees a month for power users \u2014 but the core product is free forever. And the real revenue is simple \u2014 agrochemical companies spend four thousand crore a year on TV ads hoping the right farmer sees it. We show them: this farmer, this disease, right now. Three to five rupees per recommendation. Plus API licensing for disease data, and e-commerce referrals on every treatment kit. We\u2019re not selling pesticides \u2014 we\u2019re selling the prescription.',
    ['free forever', 'four thousand crore', 'this farmer, this disease, right now', 'the prescription']
)

add_spoken(
    'The AI runs on the phone, so we barely need servers. Every single competitor needs internet. We don\u2019t.',
    ['We don\u2019t']
)

add_direction('Beat.')

add_spoken(
    'With the prize money \u2014 250 diseases. And we pilot it back home, in Assam.',
    ['250 diseases', 'Assam']
)

add_divider()

add_section_header('CLOSE [1:55\u20142:00] \u2013 Both', GREEN)

add_direction('Both stand together. Quiet.')

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.4)
r = p.add_run('GARVIT: ')
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = GREEN
r = p.add_run('KropScan.')
r.font.size = Pt(14)
r.bold = True

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.4)
r = p.add_run('ZAVIAA: ')
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = MAGENTA
r = p.add_run('Offline-first. India-first.')
r.font.size = Pt(14)
r.bold = True

add_direction('Silence. They sit down.')

doc.save('docs/KropScan_Pitch_Script_Clean.docx')
print('Formatted pitch docx saved')
