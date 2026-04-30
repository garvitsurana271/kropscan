import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor

GREEN = RGBColor(22, 101, 52)
PINK = RGBColor(157, 23, 77)
ORANGE = RGBColor(146, 64, 14)

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(0.8)
    s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(1)
    s.right_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(8)

t = doc.add_heading('KropScan \u2014 Final Pitch', level=1)
for r in t.runs:
    r.font.color.rgb = GREEN

sub = doc.add_paragraph()
r = sub.add_run('Case 7 | Garvit Surana & Zaviaa')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(100, 100, 100)
doc.add_paragraph()

with open('docs/PITCH_FINAL.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

for line in lines:
    line = line.strip()
    if not line or line.startswith('#') or line == '---':
        continue

    gm = re.match(r'\*\*GARVIT:\*\*\s*["\u201c](.+?)["\u201d]', line)
    zm = re.match(r'\*\*ZAVIAA:\*\*\s*["\u201c](.+?)["\u201d]', line)
    dm = re.match(r'^\*([^*]+)\*$', line)
    hm = re.match(r'^## (.+)', line)

    if hm:
        p = doc.add_paragraph()
        r = p.add_run(hm.group(1))
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = GREEN
    elif gm:
        p = doc.add_paragraph()
        r = p.add_run('GARVIT: ')
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = GREEN
        r = p.add_run(gm.group(1))
        r.font.size = Pt(12)
    elif zm:
        p = doc.add_paragraph()
        r = p.add_run('ZAVIAA: ')
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = PINK
        r = p.add_run(zm.group(1))
        r.font.size = Pt(12)
    elif dm:
        p = doc.add_paragraph()
        r = p.add_run(dm.group(1))
        r.italic = True
        r.font.color.rgb = ORANGE
        r.font.size = Pt(10)

doc.save('docs/KropScan_Pitch_Script_Clean.docx')
print('Clean docx saved')
