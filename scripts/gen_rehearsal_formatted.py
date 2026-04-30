"""Generate rehearsal guide directly from PITCH_FINAL.md with staging and delivery notes"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(0.8)
    s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(1)
    s.right_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

GREEN = RGBColor(56, 118, 29)
MAGENTA = RGBColor(194, 24, 91)
DARK = RGBColor(0, 0, 0)
ORANGE = RGBColor(180, 95, 6)
RED = RGBColor(204, 0, 0)
BLUE = RGBColor(17, 85, 204)
GRAY = RGBColor(120, 120, 120)


def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)


def header(text, color):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(13)
    r.font.color.rgb = color
    r.bold = True
    p.paragraph_format.space_before = Pt(6)


def stage(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run('\u25b6 ')
    r.font.color.rgb = BLUE
    r.font.size = Pt(10)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = BLUE
    r.italic = True
    p.paragraph_format.space_after = Pt(2)


def spoken(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.color.rgb = DARK


def tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run('\u26a0 ')
    r.font.size = Pt(9)
    r.font.color.rgb = ORANGE
    r.bold = True
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = ORANGE
    p.paragraph_format.space_after = Pt(2)


def kills(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run('\u2714 KILLS: ')
    r.font.size = Pt(9)
    r.font.color.rgb = GREEN
    r.bold = True
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = GREEN
    r.italic = True
    p.paragraph_format.space_after = Pt(4)


# ═══ TITLE ═══

t = doc.add_heading('KropScan \u2014 Pitch Rehearsal Guide', level=1)
for r in t.runs:
    r.font.color.rgb = GREEN

p = doc.add_paragraph()
r = p.add_run('Staging, delivery, and Q&A preempts \u2014 synced with PITCH_FINAL.md')
r.font.size = Pt(10)
r.font.color.rgb = GRAY
r.italic = True

p = doc.add_paragraph()
r = p.add_run('SETUP: ')
r.bold = True
r.font.size = Pt(10)
r.font.color.rgb = RED
r = p.add_run('Laptop + HDMI to projector. KropScan open, logged in. Model cached (do one scan with WiFi before pitch). wheat_brown_rust.jpg ready. All tabs closed. Browser fullscreen (F11). Garvit at laptop. Zaviaa slightly behind and to the right.')
r.font.size = Pt(10)

divider()

# ═══ Now read PITCH_FINAL.md and add staging notes around each section ═══

# BEAT 1: DEMO
header('DEMO [0:00\u20140:15] (15 secs) \u2013 Garvit', GREEN)

stage('Garvit at laptop. Hands on keyboard. Looks at judges first, not screen.')
stage('Click airplane mode ON \u2014 WiFi icon disappears on projected screen. Audience sees this.')

spoken('Can I diagnose a crop disease with no internet? Let\u2019s find out.')

stage('Upload wheat_brown_rust.jpg. The scanning animation plays on the big screen. ~3 seconds. CHIME.')

tip('Do NOT speak during the 3-second scan. Silence here = drama. Let the chime land.')

spoken('Wheat Brown Rust. Propiconazole. Eight hundred rupees per acre. And it only took three seconds. That too without any internet. Everything you just saw ran on this laptop \u2014 no server, no cloud, no signal.')

tip('After "no signal" \u2014 pause 2 full seconds. Shift to quieter tone for next beat.')

kills('Does it work? Offline? Treatment? Cost?')

divider()

# BEAT 2: PROBLEM
header('THE PROBLEM [0:15\u20140:35] (20 secs) \u2013 Garvit', GREEN)

stage('Step slightly back from laptop. Softer voice. This beat is about the farmer, not the app.')

spoken('But a farmer in India waits five days for that same answer \u2014 and that too, if they can find an expert at all. One agro officer for every eleven hundred farmers. Ninety thousand crore lost to disease every single year. Not because it can\u2019t be treated \u2014 because farmers can\u2019t get a diagnosis in time.')

tip('Slow down on "Ninety thousand crore." Let the number land.')

kills('Why does this matter? Problem size?')

divider()

# BEAT 3: WHY US
header('WHY US [0:35\u20140:50] (15 secs) \u2013 Zaviaa', MAGENTA)

stage('Zaviaa steps forward to stand beside Garvit. Faces judges directly. Garvit steps back \u2014 this is her moment.')

spoken('I\u2019m Zaviaa. This is Garvit. We\u2019re from Assam, where farming is all around us. And we\u2019ve watched this happen, every monsoon, every season. A farmer spots a problem, goes to a dealer, gets the wrong advice, and the crop dies. We were done watching.')

tip('"We were done watching" is the transition. Pause 1 second. Then Garvit turns to screen and takes over. Make this handoff feel natural.')

kills('Why you? Just a hackathon project?')

divider()

# BEAT 4: TECH
header('THE TECH [0:50\u20141:20] (30 secs) \u2013 Garvit', GREEN)

stage('Garvit turns back to laptop/screen. Confident pace. Gesture at screen when referencing the app.')

spoken('So we built KropScan. Every app out there \u2014 Plantix, CropIn \u2014 needs internet. Forty percent of rural India doesn\u2019t have that. So we compiled Google\u2019s EfficientNetV2 \u2014 a full neural network \u2014 to run inside the browser using WebAssembly. No server. No GPU. Just the phone. That\u2019s what you just saw.')

tip('"No server. No GPU. Just the phone." \u2014 three punches. Brief pause between each.')

spoken('And when there IS internet? The AI goes one step further \u2014 if it\u2019s not fully confident, it quietly verifies with Gemma 3 cloud. The farmer just gets the right answer \u2014 and the badge on top tells them exactly how.')

stage('Point at the result on projected screen.')

spoken('Trained on almost 2 lakh real field images \u2014 20 crops, 156 diseases.')

kills('What model? How many classes? Lab vs real? What if wrong? Plantix? Cloud contradiction?')

divider()

# BEAT 5: LANGUAGES + FEATURES
header('LANGUAGES + FEATURES [1:20\u20141:32] (12 secs) \u2013 Garvit', GREEN)

stage('Turn off airplane mode \u2014 WiFi comes back. Click Hindi in language switcher. UI changes on screen. Navigate briefly: show market prices, chatbot.')

spoken('Along with this \u2014 eleven languages, one-tap treatment purchase, plus live mandi prices, weather, crop planner, and an AI chatbot. Everything a farmer needs, in one place.')

tip('Navigate the app WHILE talking. Show each feature for ~1 second as you name it.')

kills('Languages? Can farmers act on it? Just a scanner?')

divider()

# BEAT 6: BUSINESS & SCALE
header('BUSINESS & SCALE [1:32\u20141:55] (23 secs) \u2013 Zaviaa', MAGENTA)

stage('Zaviaa steps back in, center. Garvit steps back to laptop.')

spoken('Now this is live and deployed \u2014 you can scan a leaf on your phone after we sit down.')

stage('Beat. Let that sink in.')

spoken('Free for farmers \u2014 full diagnosis, treatment, 11 languages.')

stage('Beat.')

spoken('A farmer should never see a lock icon over their medicine.')

tip('This line gets its OWN moment. Pause before and after. It\u2019s the emotional core of the business model.')

spoken('There\u2019s a Pro plan at 49 rupees a month for power users \u2014 but the core product is free forever. And the real revenue is simple \u2014 agrochemical companies spend four thousand crore a year on TV ads hoping the right farmer sees it. We show them: this farmer, this disease, right now. Three to five rupees per recommendation. Plus API licensing for disease data, and e-commerce referrals on every treatment kit. We\u2019re not selling pesticides \u2014 we\u2019re selling the prescription.')

tip('"the prescription" \u2014 say it casually. Afterthought, not mic drop.')

spoken('The AI runs on the phone, so we barely need servers. Every single competitor needs internet. We don\u2019t.')

tip('"We don\u2019t." \u2014 PAUSE before these two words.')

stage('Beat. Breath.')

spoken('With the prize money \u2014 250 diseases. And we pilot it back home, in Assam.')

tip('"Back home, in Assam" \u2014 land this softly. It\u2019s the emotional callback.')

kills('Prototype? Paywall? Business model? Revenue? Scale? Servers? Competition? Prize money?')

divider()

# BEAT 7: CLOSE
header('CLOSE [1:55\u20142:00] \u2013 Both', GREEN)

stage('Garvit steps forward beside Zaviaa. Both face judges. No fidgeting. Hands at sides.')

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
r = p.add_run('GARVIT: ')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = GREEN
r = p.add_run('KropScan.')
r.font.size = Pt(14)
r.bold = True

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
r = p.add_run('ZAVIAA: ')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = MAGENTA
r = p.add_run('Offline-first. India-first.')
r.font.size = Pt(14)
r.bold = True

stage('SILENCE. 3 seconds eye contact with judges. Walk back together. Sit down. Do NOT say thank you. Do NOT say any questions.')

tip('Practice the physical exit. Walk back together, sit at the same time.')

divider()

# ═══ STAGE MAP ═══
header('STAGE POSITIONS', BLUE)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
r = p.add_run(
    '  [JUDGES TABLE]         [PROJECTED SCREEN]\n\n'
    '                              [LAPTOP]\n'
    '                           Garvit starts here\n\n'
    '          Zaviaa starts here\n'
    '          (slightly behind, to the right)\n\n'
    '  BEAT 1-2: Garvit at laptop, Zaviaa behind\n'
    '  BEAT 3:   Zaviaa steps forward. Garvit steps back.\n'
    '  BEAT 4-5: Garvit back at laptop. Zaviaa steps back.\n'
    '  BEAT 6:   Zaviaa center. Garvit at laptop.\n'
    '  BEAT 7:   Both step forward together, center.\n'
)
r.font.size = Pt(10)
r.font.name = 'Courier New'

divider()

# ═══ PRE-PITCH CHECKLIST ═══
header('PRE-PITCH CHECKLIST', RED)

checks = [
    'Laptop charged + charger plugged in',
    'HDMI connected, screen mirrored (not extended)',
    'KropScan open in browser, logged in as Garvit Surana PREMIUM',
    'Do ONE scan with WiFi to cache the 81MB model',
    'Verify airplane mode scan works (test scan)',
    'wheat_brown_rust.jpg in Downloads or Desktop',
    'Close ALL other tabs, notifications, chat apps',
    'Browser fullscreen (F11)',
    'Language set to English (switch to Hindi during pitch)',
    'Phone on silent',
    'KropScan_QR.png open in separate tab for Q&A',
]

for c in checks:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run('\u2610 ')
    r.font.size = Pt(11)
    r = p.add_run(c)
    r.font.size = Pt(10)

doc.save('docs/KropScan_Pitch_Rehearsal.docx')
print('Rehearsal guide saved \u2014 synced with PITCH_FINAL.md')
